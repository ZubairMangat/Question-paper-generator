import streamlit as st
import pandas as pd
import PyPDF2
import docx
from docx import Document
from docx.shared import Inches
import io
import re
import random
from typing import List, Dict, Tuple
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter
import json
from datetime import datetime
from groq import Groq

# Explicitly download required NLTK data at the start
try:
    # Download and verify NLTK resources
    nltk.download('punkt', quiet=True)
    nltk.download('punkt_tab', quiet=True)
    nltk.download('stopwords', quiet=True)
    
    # Verify resources are available
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('tokenizers/punkt_tab')
    nltk.data.find('corpora/stopwords')
except LookupError as e:
    st.error(f"Failed to download NLTK resources: {str(e)}")
    st.info("Retrying NLTK resource download...")
    nltk.download('punkt', quiet=True)
    nltk.download('punkt_tab', quiet=True)
    nltk.download('stopwords', quiet=True)

class QuestionPaperGenerator:
    def __init__(self, groq_api_key=None):
        self.stop_words = set(stopwords.words('english'))
        self.groq_client = None
        if groq_api_key:
            self.groq_client = Groq(api_key=groq_api_key)
        
    def extract_text_from_pdf(self, file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def extract_template_from_docx(self, file) -> str:
        """Extract template structure from DOCX file"""
        try:
            doc = docx.Document(file)
            template = ""
            for paragraph in doc.paragraphs:
                template += paragraph.text + "\n"
            return template
        except Exception as e:
            st.error(f"Error reading template: {str(e)}")
            return ""
    
    def find_topic_content(self, full_text: str, topic: str) -> str:
        """Extract content related to specific topic/chapter"""
        lines = full_text.split('\n')
        topic_content = []
        capturing = False
        
        topic_patterns = [
            rf"chapter\s+\d*\s*:?\s*{re.escape(topic)}",
            rf"unit\s+\d*\s*:?\s*{re.escape(topic)}",
            rf"section\s+\d*\s*:?\s*{re.escape(topic)}",
            rf"{re.escape(topic)}",
        ]
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            topic_lower = topic.lower()
            
            if any(re.search(pattern, line_lower, re.IGNORECASE) for pattern in topic_patterns):
                capturing = True
                topic_content.append(line)
                continue
            
            if capturing and (
                re.search(r'chapter\s+\d+', line_lower) or 
                re.search(r'unit\s+\d+', line_lower) or
                re.search(r'section\s+\d+', line_lower)
            ) and topic_lower not in line_lower:
                break
            
            if capturing:
                topic_content.append(line)
        
        if not topic_content:
            sentences = sent_tokenize(full_text)
            return ' '.join(sentences[:50])
        
        return '\n'.join(topic_content)
    
    def generate_mcqs(self, content: str, num_questions: int) -> List[Dict]:
        """Generate Multiple Choice Questions using Groq API with fallback"""
        if not self.groq_client:
            return self.generate_mcqs_fallback(content, num_questions)
        
        try:
            prompt = f"""
            Based on the following educational content, generate {num_questions} multiple choice questions (MCQs).
            
            Content:
            {content[:3000]}
            
            For each MCQ, provide:
            1. A clear question
            2. Four options (A, B, C, D)
            3. Indicate the correct answer
            
            Format your response as JSON with this structure:
            {{
                "mcqs": [
                    {{
                        "question": "Question text here?",
                        "options": ["Option A", "Option B", "Option C", "Option D"],
                        "correct_answer": "Option A"
                    }}
                ]
            }}
            
            Make questions educational and relevant to the content.
            """
            
            response = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are an expert educator who creates high-quality multiple choice questions."},
                    {"role": "user", "content": prompt}
                ],
                model="llama3-8b-8192",
                temperature=0.7,
                max_tokens=2000
            )
            
            response_text = response.choices[0].message.content
            
            try:
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                if start_idx != -1 and end_idx != -1:
                    json_str = response_text[start_idx:end_idx]
                    data = json.loads(json_str)
                    return data.get('mcqs', [])
            except:
                pass
            
            return self.parse_mcqs_from_text(response_text, num_questions)
            
        except Exception as e:
            st.warning(f"Groq API error: {str(e)}. Using fallback method.")
            return self.generate_mcqs_fallback(content, num_questions)
    
    def generate_short_questions(self, content: str, num_questions: int) -> List[str]:
        """Generate short answer questions using Groq API with fallback"""
        if not self.groq_client:
            return self.generate_short_questions_fallback(content, num_questions)
        
        try:
            prompt = f"""
            Based on the following educational content, generate {num_questions} short answer questions.
            
            Content:
            {content[:3000]}
            
            Create questions that:
            1. Can be answered in 2-3 sentences
            2. Test understanding of key concepts
            3. Are clear and specific
            
            Format: Just list the questions, one per line, numbered.
            
            Example:
            1. What is the main principle of...?
            2. Define the term...
            3. Explain briefly how...
            """
            
            response = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are an expert educator who creates concise, focused questions."},
                    {"role": "user", "content": prompt}
                ],
                model="llama3-8b-8192",
                temperature=0.7,
                max_tokens=1000
            )
            
            response_text = response.choices[0].message.content
            questions = []
            
            for line in response_text.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    question = re.sub(r'^\d+\.?\s*', '', line)
                    question = re.sub(r'^[-•]\s*', '', question)
                    if question:
                        questions.append(question)
            
            return questions[:num_questions]
            
        except Exception as e:
            st.warning(f"Groq API error: {str(e)}. Using fallback method.")
            return self.generate_short_questions_fallback(content, num_questions)
    
    def generate_long_questions(self, content: str, num_questions: int) -> List[str]:
        """Generate long answer questions using Groq API with fallback"""
        if not self.groq_client:
            return self.generate_long_questions_fallback(content, num_questions)
        
        try:
            prompt = f"""
            Based on the following educational content, generate {num_questions} long answer questions.
            
            Content:
            {content[:3000]}
            
            Create questions that:
            1. Require detailed explanations (5-10 sentences)
            2. Test deep understanding and analysis
            3. Encourage critical thinking
            4. May ask for examples, comparisons, or evaluations
            
            Use question starters like:
            - "Discuss in detail..."
            - "Analyze and explain..."
            - "Compare and contrast..."
            - "Evaluate the importance of..."
            - "Examine the relationship between..."
            
            Format: Just list the questions, one per line, numbered.
            """
            
            response = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are an expert educator who creates comprehensive analytical questions."},
                    {"role": "user", "content": prompt}
                ],
                model="llama3-8b-8192",
                temperature=0.8,
                max_tokens=1000
            )
            
            response_text = response.choices[0].message.content
            questions = []
            
            for line in response_text.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    question = re.sub(r'^\d+\.?\s*', '', line)
                    question = re.sub(r'^[-•]\s*', '', question)
                    if question:
                        questions.append(question)
            
            return questions[:num_questions]
            
        except Exception as e:
            st.warning(f"Groq API error: {str(e)}. Using fallback method.")
            return self.generate_long_questions_fallback(content, num_questions)
    
    def parse_mcqs_from_text(self, text: str, num_questions: int) -> List[Dict]:
        """Parse MCQs from Groq response text"""
        mcqs = []
        lines = text.split('\n')
        current_mcq = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.endswith('?') and len(line) > 20:
                if current_mcq and 'question' in current_mcq:
                    mcqs.append(current_mcq)
                current_mcq = {'question': line, 'options': [], 'correct_answer': '', 'context': line[:100] + "..." if len(line) > 100 else line}
            
            elif re.match(r'^[A-D][\.\)]\s*', line):
                if current_mcq:
                    option_text = re.sub(r'^[A-D][\.\)]\s*', '', line)
                    current_mcq['options'].append(option_text)
            
            elif 'correct' in line.lower() or 'answer' in line.lower():
                if current_mcq and current_mcq['options']:
                    for option in current_mcq['options']:
                        if option.lower() in line.lower():
                            current_mcq['correct_answer'] = option
                            break
        
        if current_mcq and 'question' in current_mcq:
            mcqs.append(current_mcq)
        
        return mcqs[:num_questions]
    
    def generate_mcqs_fallback(self, content: str, num_questions: int) -> List[Dict]:
        """Fallback MCQ generation without API"""
        sentences = sent_tokenize(content)
        mcqs = []
        
        for i in range(min(num_questions, len(sentences))):
            sentence = sentences[i].strip()
            if len(sentence) < 20:
                continue
            
            words = word_tokenize(sentence.lower())
            words = [w for w in words if w.isalnum() and w not in self.stop_words]
            
            if len(words) < 3:
                continue
            
            key_word = random.choice(words[:min(5, len(words))])
            
            question_text = sentence.replace(key_word, "______", 1)
            if question_text == sentence:
                question_text = f"What is mentioned about {key_word} in the following context?"
            
            correct_answer = key_word.title()
            wrong_answers = self.generate_wrong_answers(correct_answer, words)
            
            options = [correct_answer] + wrong_answers[:3]
            random.shuffle(options)
            
            mcq = {
                'question': question_text,
                'options': options,
                'correct_answer': correct_answer,
                'context': sentence[:100] + "..." if len(sentence) > 100 else sentence
            }
            mcqs.append(mcq)
        
        return mcqs[:num_questions]
    
    def generate_short_questions_fallback(self, content: str, num_questions: int) -> List[str]:
        """Fallback short question generation"""
        sentences = sent_tokenize(content)
        questions = []
        
        question_starters = [
            "What is",
            "Define",
            "Explain briefly",
            "What are the main",
            "How does",
            "Why is",
            "What causes",
            "List the",
        ]
        
        for i in range(min(num_questions, len(sentences))):
            sentence = sentences[i].strip()
            if len(sentence) < 20:
                continue
            
            words = word_tokenize(sentence)
            words = [w for w in words if w.isalnum() and len(w) > 3]
            
            if not words:
                continue
            
            key_concept = random.choice(words[:min(3, len(words))])
            starter = random.choice(question_starters)
            
            if starter in ["Define", "List the"]:
                question = f"{starter} {key_concept}."
            else:
                question = f"{starter} {key_concept}?"
            
            questions.append(question)
        
        return questions[:num_questions]
    
    def generate_long_questions_fallback(self, content: str, num_questions: int) -> List[str]:
        """Fallback long question generation"""
        sentences = sent_tokenize(content)
        questions = []
        
        question_templates = [
            "Discuss in detail about {}.",
            "Explain the concept of {} with examples.",
            "Analyze the importance of {} in the given context.",
            "Compare and contrast different aspects of {}.",
            "Evaluate the role of {} and its implications.",
            "Describe the process of {} step by step.",
            "Examine the relationship between {} and related concepts.",
        ]
        
        key_phrases = []
        for sentence in sentences[:20]:
            words = word_tokenize(sentence)
            words = [w for w in words if w.isalnum() and len(w) > 3 and w.lower() not in self.stop_words]
            
            for i in range(len(words) - 1):
                phrase = f"{words[i]} {words[i+1]}"
                key_phrases.append(phrase)
        
        phrase_counts = Counter(key_phrases)
        common_phrases = [phrase for phrase, count in phrase_counts.most_common(10)]
        
        for i in range(min(num_questions, len(common_phrases))):
            template = random.choice(question_templates)
            phrase = common_phrases[i]
            question = template.format(phrase)
            questions.append(question)
        
        return questions[:num_questions]
    
    def generate_wrong_answers(self, correct_answer: str, word_pool: List[str]) -> List[str]:
        """Generate plausible wrong answers"""
        wrong_answers = []
        
        candidates = [w.title() for w in word_pool if w != correct_answer.lower()]
        wrong_answers.extend(random.sample(candidates, min(2, len(candidates))))
        
        generic_answers = ["None of the above", "All of the above", "Not specified", "Cannot be determined"]
        wrong_answers.extend(random.sample(generic_answers, 2))
        
        return wrong_answers[:3]
    
    def format_question_paper(self, mcqs: List[Dict], short_questions: List[str], 
                            long_questions: List[str], template: str = None) -> str:
        """Format the generated questions into a question paper"""
        
        if template and template.strip():
            paper = template + "\n\n"
        else:
            paper = f"""
QUESTION PAPER
==============

Subject: [Subject Name]
Date: {datetime.now().strftime('%Y-%m-%d')}
Time: [Duration]
Maximum Marks: [Total Marks]

Instructions:
1. Read all questions carefully
2. Answer all questions
3. Write clearly and legibly

"""
        
        if mcqs:
            paper += "\nSECTION A: MULTIPLE CHOICE QUESTIONS\n"
            paper += "=" * 40 + "\n\n"
            
            for i, mcq in enumerate(mcqs, 1):
                paper += f"Q{i}. {mcq['question']}\n"
                for j, option in enumerate(mcq['options']):
                    paper += f"    {chr(65+j)}. {option}\n"
                paper += "\n"
        
        if short_questions:
            paper += "\nSECTION B: SHORT ANSWER QUESTIONS\n"
            paper += "=" * 40 + "\n\n"
            
            for i, question in enumerate(short_questions, 1):
                paper += f"Q{i}. {question}\n\n"
        
        if long_questions:
            paper += "\nSECTION C: LONG ANSWER QUESTIONS\n"
            paper += "=" * 40 + "\n\n"
            
            for i, question in enumerate(long_questions, 1):
                paper += f"Q{i}. {question}\n\n"
        
        return paper
    
    def create_docx_output(self, formatted_paper: str) -> io.BytesIO:
        """Create a DOCX file from the formatted paper"""
        doc = Document()
        
        title = doc.add_heading('Question Paper', 0)
        title.alignment = 1
        
        lines = formatted_paper.split('\n')
        current_paragraph = []
        
        for line in lines:
            if line.strip() == "":
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                continue
            
            if line.startswith('SECTION') or line.startswith('Q') or line.startswith('='):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                
                if line.startswith('SECTION'):
                    doc.add_heading(line, level=1)
                elif line.startswith('Q'):
                    doc.add_paragraph(line, style='List Number')
                elif not line.startswith('='):
                    doc.add_paragraph(line)
            else:
                current_paragraph.append(line.strip())
        
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

def main():
    st.set_page_config(
        page_title="AI Question Paper Generator",
        page_icon="📝",
        layout="wide"
    )
    
    st.title("🤖 AI-Powered Question Paper Generator (Groq Powered)")
    st.markdown("Generate customized question papers from textbooks using free Groq API!")
    
    # Read Groq API key from Streamlit secrets
    try:
        groq_api_key = st.secrets["GROQ_API_KEY"]
    except KeyError:
        st.error("Groq API key not found in secrets. Please configure it in Streamlit Cloud settings.")
        return
    
    generator = QuestionPaperGenerator(groq_api_key)
    
    with st.sidebar:
        st.header("🔑 API Configuration")
        st.success("✅ Groq API Key configured!")
        
        st.header("📚 Upload Files")
        
        textbook_file = st.file_uploader(
            "Upload Textbook", 
            type=['pdf', 'docx'],
            help="Upload your textbook in PDF or DOCX format"
        )
        
        template_file = st.file_uploader(
            "Upload Question Paper Template (Optional)", 
            type=['docx', 'txt'],
            help="Upload your question paper template"
        )
        
        st.header("🎯 Topic Selection")
        topic_input = st.text_input(
            "Enter Topic/Chapter Name",
            placeholder="e.g., Introduction to Physics, Chapter 5",
            help="Enter the specific topic or chapter you want questions from"
        )
        
        st.header("📊 Question Configuration")
        
        col1, col2 = st.columns(2)
        with col1:
            num_mcqs = st.number_input("MCQs", min_value=0, max_value=20, value=5)
            num_short = st.number_input("Short Questions", min_value=0, max_value=15, value=3)
        
        with col2:
            num_long = st.number_input("Long Questions", min_value=0, max_value=10, value=2)
            total_questions = num_mcqs + num_short + num_long
            st.metric("Total Questions", total_questions)
        
        st.header("🤖 AI Model Status")
        st.success("🚀 Using Groq AI (Enhanced)")
        
        generate_button = st.button("🚀 Generate Question Paper", type="primary")
    
    if textbook_file is None:
        st.info("👈 Please upload a textbook file to get started!")
        
        st.markdown("## 🎯 How it works:")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            ### 📚 Step 1: Upload
            - Upload your textbook (PDF/DOCX)
            - Upload template (optional)
            - Specify topic/chapter
            """)
        
        with col2:
            st.markdown("""
            ### ⚙️ Step 2: Configure
            - Set number of MCQs
            - Set short questions count
            - Set long questions count
            """)
        
        with col3:
            st.markdown("""
            ### 📝 Step 3: Generate
            - AI extracts relevant content
            - Generates intelligent questions using Groq API
            - Formats according to template
            """)
        
        st.markdown("## 🌟 Why Use Groq API?")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            ### 🚀 **Enhanced Quality**
            - Smarter question generation
            - Better context understanding
            - More relevant questions
            """)
        
        with col2:
            st.markdown("""
            ### 🆓 **Completely Free**
            - No usage limits
            - Fast processing
            - Professional results
            """)
        
        with col3:
            st.markdown("""
            ### 🎯 **Intelligent Analysis**
            - Understands content deeply
            - Creates pedagogically sound questions
            - Maintains academic standards
            """)
        
        return
    
    if generate_button:
        if not topic_input.strip():
            st.error("Please enter a topic or chapter name!")
            return
        
        with st.spinner("Processing textbook and generating questions..."):
            if textbook_file.type == "application/pdf":
                full_text = generator.extract_text_from_pdf(textbook_file)
            else:
                full_text = generator.extract_text_from_docx(textbook_file)
            
            if not full_text.strip():
                st.error("Could not extract text from the uploaded file!")
                return
            
            topic_content = generator.find_topic_content(full_text, topic_input)
            
            if not topic_content.strip():
                st.warning("Could not find specific content for the topic. Using general content.")
                topic_content = full_text[:5000]
            
            template_text = ""
            if template_file:
                if template_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    template_text = generator.extract_template_from_docx(template_file)
                else:
                    template_text = str(template_file.read(), "utf-8")
            
            mcqs = generator.generate_mcqs(topic_content, num_mcqs) if num_mcqs > 0 else []
            short_questions = generator.generate_short_questions(topic_content, num_short) if num_short > 0 else []
            long_questions = generator.generate_long_questions(topic_content, num_long) if num_long > 0 else []
            
            formatted_paper = generator.format_question_paper(mcqs, short_questions, long_questions, template_text)
            
            st.session_state.formatted_paper = formatted_paper
            st.session_state.mcqs = mcqs
            st.session_state.short_questions = short_questions
            st.session_state.long_questions = long_questions
    
    if hasattr(st.session_state, 'formatted_paper'):
        st.success("✅ Question paper generated successfully!")
        
        tab1, tab2, tab3 = st.tabs(["📝 Preview", "📊 Question Analysis", "⬇️ Download"])
        
        with tab1:
            st.markdown("### Question Paper Preview")
            st.text_area(
                "Generated Question Paper",
                value=st.session_state.formatted_paper,
                height=600,
                key="preview_area"
            )
        
        with tab2:
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("MCQs Generated", len(st.session_state.mcqs))
                if st.session_state.mcqs:
                    with st.expander("View MCQ Details"):
                        for i, mcq in enumerate(st.session_state.mcqs[:3], 1):
                            st.write(f"**Q{i}:** {mcq['question'][:50]}...")
            
            with col2:
                st.metric("Short Questions", len(st.session_state.short_questions))
                if st.session_state.short_questions:
                    with st.expander("View Short Questions"):
                        for i, q in enumerate(st.session_state.short_questions[:3], 1):
                            st.write(f"**Q{i}:** {q}")
            
            with col3:
                st.metric("Long Questions", len(st.session_state.long_questions))
                if st.session_state.long_questions:
                    with st.expander("View Long Questions"):
                        for i, q in enumerate(st.session_state.long_questions[:3], 1):
                            st.write(f"**Q{i}:** {q}")
            
            st.markdown("### Quality Indicators")
            col1, col2 = st.columns(2)
            
            with col1:
                st.success("🤖 Enhanced AI Generation")
                st.info("Questions generated using advanced Groq AI model")
            
            with col2:
                total_words = len(st.session_state.formatted_paper.split())
                st.metric("Total Words", total_words)
                
                estimated_time = (len(st.session_state.mcqs) * 2) + (len(st.session_state.short_questions) * 5) + (len(st.session_state.long_questions) * 15)
                st.metric("Estimated Time", f"{estimated_time} min")
        
        with tab3:
            st.markdown("### Download Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                docx_file = generator.create_docx_output(st.session_state.formatted_paper)
                st.download_button(
                    label="📄 Download as DOCX",
                    data=docx_file.getvalue(),
                    file_name=f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                st.download_button(
                    label="📄 Download as TXT",
                    data=st.session_state.formatted_paper,
                    file_name=f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
            
            st.info("💡 Tip: The DOCX format preserves better formatting for printing!")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.info("Try restarting the app if issues persist.")
